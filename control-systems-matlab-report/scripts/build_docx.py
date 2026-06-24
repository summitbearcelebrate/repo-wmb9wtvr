# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(BASE, "figs")
OUT = os.path.join(BASE, "自动控制原理实验报告_薄天乐.docx")

doc = Document()

# ---- default styles (CJK) ----
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(10.5)
normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

def set_cjk(run, font="宋体", size=10.5, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element.rPr.rFonts
    r.set(qn("w:eastAsia"), font)

def para(text="", size=10.5, bold=False, align=None, font="宋体", space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        set_cjk(run, font, size, bold)
    return p

def heading(text, size=14, font="黑体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    set_cjk(run, font, size, True)
    return p

def sub(text, size=11, font="黑体"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_cjk(run, font, size, True)
    return p

def img(path, width=3.0, center=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    p.add_run().add_picture(os.path.join(FIG, path), width=Inches(width))
    return p

def cell_set(cell, content, size=10, bold=False, font="宋体"):
    cell.paragraphs[0].text = ""
    run = cell.paragraphs[0].add_run(content)
    set_cjk(run, font, size, bold)

def cell_img(cell, path, width=2.6, caption=None):
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(os.path.join(FIG, path), width=Inches(width))

def add_run_to_cell(cell, text, size=10, bold=False, font="宋体", newpara=True):
    if newpara:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]
    run = p.add_run(text)
    set_cjk(run, font, size, bold)
    return p

def set_fixed_layout(t, col1w, col2w):
    t.autofit = False
    t.allow_autofit = False
    tblPr = t._tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = tblPr.makeelement(qn("w:tblLayout"), {})
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    # set grid column widths
    grid = t._tbl.find(qn("w:tblGrid"))
    cols = grid.findall(qn("w:gridCol"))
    widths = [col1w, col2w]
    for gc, w in zip(cols, widths):
        gc.set(qn("w:w"), str(int(w * 1440)))

def set_cell_width(cell, w_in):
    cell.width = Inches(w_in)
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = tcPr.makeelement(qn("w:tcW"), {})
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(w_in * 1440)))
    tcW.set(qn("w:type"), "dxa")

# rows: list of (desc, [ (type,val,..), ... ]) ; build 2-col table
def make_table(header, rows, col1w=2.0, col2w=4.1):
    t = doc.add_table(rows=1, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_fixed_layout(t, col1w, col2w)
    hc = t.rows[0].cells
    cell_set(hc[0], header[0], bold=True); hc[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_set(hc[1], header[1], bold=True); hc[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for desc, items in rows:
        r = t.add_row().cells
        cell_set(r[0], desc)
        first = True
        for it in items:
            if it[0] == "text":
                add_run_to_cell(r[1], it[1], font="Consolas" if it[1].strip().startswith(">>") or "=" in it[1] else "宋体", newpara=not first)
            elif it[0] == "img":
                w = it[2] if len(it) > 2 else 2.6
                w = min(w, col2w - 0.35)
                if first:
                    p = r[1].paragraphs[0]
                else:
                    p = r[1].add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(os.path.join(FIG, it[1]), width=Inches(w))
            first = False
    # column widths (fixed)
    for row in t.rows:
        set_cell_width(row.cells[0], col1w)
        set_cell_width(row.cells[1], col2w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t

# ============================================================
# 封面
# ============================================================
para("", space_after=10)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("实验(实训)课程报告"); set_cjk(r, "黑体", 22, True)
para("", space_after=20)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("课程      自动控制原理"); set_cjk(r, "黑体", 16, True)
para("", space_after=20)

info = doc.add_table(rows=4, cols=4); info.style = "Table Grid"
info.alignment = WD_TABLE_ALIGNMENT.CENTER
data = [
    ("学    院", "计算机与人工智能学院", "专   业", "电子信息工程"),
    ("年    级", "23级", "班   级", "电信6班"),
    ("学生姓名", "薄天乐", "学   号", "202308630603"),
    ("指导教师", "姜锦云", "职   称", "讲师"),
]
for i, rowd in enumerate(data):
    cells = info.rows[i].cells
    for j, val in enumerate(rowd):
        cell_set(cells[j], val, size=11, bold=(j % 2 == 0))
        cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
para("", space_after=40)
para("湖南财政经济学院", size=16, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
para("2026年 6 月", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")

doc.add_page_break()

# ============================================================
# 实验一
# ============================================================
heading("实验一  系统的数学模型")
sub("一、实验目的")
para("1．掌握MATLAB有关传递函数求取及其零、极点计算的函数。")
para("2．掌握用MATLAB求取系统的数学模型。")
sub("二、实验内容与步骤")
para("1．准备知识", bold=True, font="黑体", size=10.5)
para("（1）求串联环节的传递函数：串联后的传递函数为两环节传递函数之积。")
para("MATLAB计算公式：[num,den]=series(num1,den1,num2,den2)", font="Consolas")
para("（2）求并联环节的传递函数：并联后的传递函数为两环节传递函数之和。")
para("MATLAB计算公式：[num,den]=parallel(num1,den1,num2,den2)", font="Consolas")
para("（3）求单位反馈控制系统的传递函数：")
para("MATLAB计算公式：[num,den]=cloop(num1,den1,sign)", font="Consolas")
para("Sign参数：正反馈用+1，负反馈用-1。缺省情况为负反馈。")
para("（4）求闭环控制系统的传递函数：")
para("MATLAB计算公式：[num,den]=feedback(num1,den1,num2,den2,sign)", font="Consolas")
para("Sign参数：正反馈用+1，负反馈用-1。缺省情况为负反馈。")
para("（5）多项式相乘需先建立两个多项式对应的向量a、b，然后利用conv()函数。例：")
para(">>a=[1,2];\n>>b=[2,3];\n>>c=conv(a,b)", font="Consolas")
para("上面三个命令就是求取多项式(s+2)与(2s+3)相乘后的向量。")

para("2．特征多项式的建立与特征根的求取", bold=True, font="黑体", size=10.5)
para("表1-1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令"), [
    ("构建特征多项式p(s)=s³+3s²+4的矩阵", [("text", ">>p=[1,3,0,4];")]),
    ("求特征方程p(s)=s³+3s²+4=0的特征根", [("text", ">>r=roots(p)"), ("img", "e1_t11_roots.png", 2.4)]),
    ("从特征根构建特征多项式的矩阵", [("text", ">>p=poly(r)"), ("img", "e1_t11_poly.png", 3.2)]),
])

para("3．求单位反馈系统的传递函数：", bold=True, font="黑体", size=10.5)
make_table(("说明", "命令"), [
    ("构建传递函数 G(s)=1/(500s²) 的特征多项式", [("text", ">>numg=[1];\n>>deng=[500,0,0];")]),
    ("构建传递函数 Gc(s)=(s+1)/(s+2) 的特征多项式", [("text", ">>numc=[1,1];\n>>denc=[1,2];")]),
    ("求 G(s)·Gc(s)", [("text", ">>[num1,den1]=series(numg,deng,numc,denc);"), ("img", "e1_t11_series.png", 3.2)]),
    ("求开环传递函数的闭环传递函数（单位负反馈）", [("text", ">>[num,den]=cloop(num1,den1);\nnum=[1 1]  den=[500 1000 1 1]")]),
    ("输出传递函数", [("text", ">>printsys(num,den);"), ("img", "e1_t11_cloop.png", 3.2)]),
])

para("3．传递函数零、极点的求取", bold=True, font="黑体", size=10.5)
para("表1-2", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令"), [
    ("构建传递函数 G(s)=(s²+3s+2)/(s³+3s²+4s+12)", [("text", ">>num1=[1,3,2];den1=[1,3,4,12];")]),
    ("求G(s)的零点", [("text", ">>z=roots(num1);"), ("img", "e1_t12_zeros.png", 2.2)]),
    ("求G(s)的极点", [("text", ">>p=roots(den1);"), ("img", "e1_t12_poles.png", 2.4)]),
    ("定义因子", [("text", ">>n1=[1,1];n2=[1,2];d1=[1,2*i];d2=[1,-2*i];d3=[1,3];")]),
    ("求多项式(s+1)(s+2)", [("text", "num2=conv(n1,n2)"), ("img", "e1_t12_num2.png", 2.6)]),
    ("求多项式(s-2j)(s+2j)(s+3)", [("text", "den2=conv(conv(d1,d2),d3)"), ("img", "e1_t12_den2.png", 2.8)]),
    ("构建并输出结果", [("text", "printsys(num2,den2)"), ("img", "e1_t12_printsys2.png", 3.2)]),
    ("构建特征多项式矩阵", [("text", ">>num=conv(num1,den2);den=conv(den1,num2);")]),
    ("输出以多项式表示的传递函数", [("text", ">>printsys(num,den)"), ("img", "e1_t12_printsys.png", 3.2)]),
    ("输出传递函数的极点和零点图", [("text", ">>pzmap(num, den);\n>>title('极点一零点图');"), ("img", "e1_t12_pzmap.png", 3.0)]),
])

para("3．求反馈系统的传递函数：", bold=True, font="黑体", size=10.5)
para("表1-3", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令"), [
    ("构建传递函数 G(s)=1/(500s²) 的特征多项式", [("text", ">>numg=[1];\n>>deng=[500,0,0];")]),
    ("构建传递函数 H(s)=(s+1)/(s+2) 的特征多项式", [("text", ">>numh=[1,1];\n>>denh=[1,2];")]),
    ("求闭环传递函数（负反馈）", [("text", ">>[num,den]=feedback(numg,deng,numh,denh);")]),
    ("输出传递函数", [("text", ">>printsys(num,den)"), ("img", "e1_t13_feedback.png", 3.2)]),
])

sub("三、实验报告")
para("1．完成表1-1至1-3（见上）。")
para("2．某特征方程的特征根分别为3，5，2，试求其特征方程。")
para("解：>>p=poly([3 5 2]) → p = [1  -10  31  -30]，即特征方程为 s³-10s²+31s-30=0。", font="Consolas")
para("3．利用MATLAB命令求取以下系统传递函数及零极点图，并记录下结果。")
para("题给系统为方框图（嵌套反馈）：前向通道 G1(s)=10/(s+1) 与 G2(s)=2/[s(s+1)] 串联，内环经 H1(s)=(s+2)/(s+3) 正反馈，外环经 H2(s)=5s/(s²+6s+8) 负反馈。")
para("解题命令：", bold=True)
para(">>s=tf('s');\n>>G1=10/(s+1); G2=2/(s*(s+1));\n>>H1=(s+2)/(s+3); H2=5*s/(s^2+6*s+8);\n>>M=feedback(G2,H1,+1);       % 内环正反馈\n>>T=feedback(G1*M,H2,-1);     % 外环负反馈\n>>T=minreal(T)\n>>zero(T), pole(T)\n>>pzmap(T); title('极点一零点图');", font="Consolas")
para("闭环传递函数 T(s)：", bold=True)
img("e1_hw3_tf.png", width=4.2)
para("零点、极点：", bold=True)
img("e1_hw3_zp.png", width=3.0)
para("极点—零点图：", bold=True)
img("e1_hw3_pzmap.png", width=3.2)
para("结论：系统有 2 个具有正实部的极点（约 0.43±1.88j 及 0.12），位于 s 右半平面，故该闭环系统不稳定。")

doc.add_page_break()

# ============================================================
# 实验二  控制系统的时域分析
# ============================================================
heading("实验二  控制系统的时域分析")
sub("一、实验目的")
para("1．掌握用MATLAB对系统进行时间响应分析。")
para("2．掌握一阶惯性系统以及二阶系统的时间响应特征以及系统性能与系统参数之间的关系。")
sub("二、实验内容与步骤")
para("1．准备知识", bold=True, font="黑体", size=10.5)
para("设输入x(t)，输出为y(t)，仿真时间段为矩阵t。利用MATLAB求取系统时间响应的函数有：")
para("求取单位阶跃响应：step(num,den)", font="Consolas")
para("求取单位脉冲响应：impulse(num,den)", font="Consolas")
para("求取任意输入的时间响应：lsim(num,den,u,t)（u表示输入列向量）", font="Consolas")

para("2．按要求填表。", bold=True, font="黑体", size=10.5)
para("表2-1  求一阶系统的单位阶跃响应曲线", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令或结果"), [
    ("用MATLAB求一阶惯性系统的单位阶跃响应曲线 y(t)=10(1-e^(-2t))，依次运行下列命令并记录结果",
     [("text", ">>t=[0:.5:5];\n>>y=10*(1-exp(-2*t));\n>>plot(t,y,'r');\n>>axis([0 5 0 10.1]);\n>>title('y(t)=1-exp(-2t)');\n>>xlabel('t');ylabel('y(t)');grid on"),
      ("img", "e2_t21_manual.png", 3.0)]),
    ("用step函数实现上一行的一阶惯性系统 10/(0.5s+1) 的单位阶跃响应曲线",
     [("text", "num=[10]; den=[0.5,1];\nstep(num,den); grid on;\ntitle('Step Response of 10/(0.5s+1)');"),
      ("img", "e2_t21_step.png", 3.0)]),
    ("将 10/(0.5s+1) 与 10/(s+1) 绘于同一张图，不同颜色并加图例，比较结论",
     [("text", "num1=[10];den1=[0.5,1];\nnum2=[10];den2=[1,1];\nstep(num1,den1);hold on;step(num2,den2);\nlegend('0.5s+1','s+1');grid on;title('Comparison');"),
      ("img", "e2_t21_compare.png", 3.0),
      ("text", "结论：时间常数越小（0.5 vs 1），响应越快，到达稳态值的时间越短。")]),
])

para("表2-2  求二阶系统的单位阶跃响应曲线", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令或结果"), [
    ("用MATLAB求二阶系统 ωn²/(s²+2ζωn·s+ωn²) 的单位阶跃响应曲线（ωn=0.4，ζ从0变化到2），记录结果",
     [("text", "t = linspace(0, 80, 1000); wn = 0.4;\nzetas = [0:0.2:0.8,1:0.5:2];\nfigure; hold on; grid on;\nfor zeta = zetas\n  sys = tf(wn^2, [1, 2*zeta*wn, wn^2]);\n  [y,t_out] = step(sys, t);\n  plot(t_out,y,'DisplayName',['\\zeta = ',num2str(zeta)]);\nend\ntitle('二阶系统响应曲线');\nxlabel('时间 (s)');ylabel('响应幅值');\nlegend('show');axis([0 80 0 2]);"),
      ("img", "e2_t22_second.png", 3.2),
      ("text", "结论：\n-ζ=0时等幅振荡；\n-0<ζ<1时欠阻尼，超调量随ζ增大而减小，调节时间先减后增；\n-ζ=1临界阻尼，无超调；\n-ζ>1过阻尼，响应缓慢，无振荡。")]),
])

para("表2-3  求系统脉冲响应", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "结果"), [
    ("试求系统 G(s)=1/(s²+0.2s+1) 的单位脉冲响应",
     [("text", ">>num=[0 0 1]; den=[1 0.2 1];\n>>impulse(num,den); grid on;\n>>title('Unit-impulse Response of G(s)=1/(s^2+0.2s+1)')"),
      ("img", "e2_t23_impulse.png", 3.0),
      ("text", "得到衰减振荡曲线，超调量约50%，振荡频率约1 rad/s。")]),
])

sub("三、实验报告")
para("1．完成表2-1至2-3（见上）。")
para("2．观察step()的调用格式，绘制系统 G(s)=(s²+3s+7)/(s⁴+4s³+6s²+4s+1) 的阶跃响应曲线。")
para(">>num=[1 3 7]; den=[1 4 6 4 1];\n>>sys=tf(num,den);\n>>step(sys); grid on;\n>>title('Step Response');", font="Consolas")
img("e2_hw2_step.png", width=3.6)
para("3．对典型二阶系统 ωn²/(s²+2ζωn·s+ωn²)：")
para("1）ωn 一定（取 ωn=1），ζ 取 0,0.25,0.5,1.0,2.0 时的单位阶跃响应曲线：")
para(">>t=linspace(0,20,1500);\n>>for z=[0 0.25 0.5 1.0 2.0]\n>>  sys=tf(1,[1 2*z 1]); step(sys,t); hold on;\n>>end", font="Consolas")
img("e2_hw3_zeta.png", width=3.6)
para("分析：ζ 越大，超调量越小、振荡越弱；ζ=0 等幅振荡，0<ζ<1 欠阻尼，ζ=1 临界阻尼无超调，ζ>1 过阻尼无振荡。")
para("2）ζ=0.25、ωn 取 1,2,4,6 时的单位阶跃响应曲线：")
para(">>t=linspace(0,20,1500);\n>>for wn=[1 2 4 6]\n>>  sys=tf(wn^2,[1 2*0.25*wn wn^2]); step(sys,t); hold on;\n>>end", font="Consolas")
img("e2_hw3_wn.png", width=3.6)
para("分析：ωn 越大，上升时间、峰值时间和调节时间越短（响应越快），而超调量仅由 ζ 决定，基本不变。")

doc.add_page_break()

# ============================================================
# 实验三
# ============================================================
heading("实验三  系统稳定性判断、频域分析及根轨迹绘制")
sub("一、实验目的")
para("1．熟练掌握系统稳定性的判断方法；")
para("2．利用MATLAB画一阶和二阶控制系统的伯德图；")
para("3．利用MATLAB绘制线性系统的根轨迹；")
para("4．利用MATLAB计算所给系统的相角裕量和幅值裕量。")
sub("二、实验内容与步骤")
para("1．准备知识", bold=True, font="黑体", size=10.5)
para("(1) 系统稳定性判断：①直接求根判稳 roots()，系统稳定的充要条件是特征方程的根均具有负实部；②劳斯稳定判据（需自定义routh函数）。")
para("(2) 频域分析：nyquist(num,den) 绘制奈奎斯特图；bode(num,den) 绘制波特图，幅值 magdb=20·log10(mag)。")
para("(3) 根轨迹：rlocus(num,den) 绘制根轨迹，rlocus(num,den,k) 可人工设定增益范围。")

para("2．按要求填表。", bold=True, font="黑体", size=10.5)
para("表3-1  系统稳定性判断", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令或结果"), [
    ("特征方程 s⁴+3s³+2s²+5s+1=0，求根并判断稳定性",
     [("text", "den=[1,3,2,5,1];\nroots(den)"), ("img", "e3_t31_roots.png", 2.6),
      ("text", "该系统不稳定，因为有两个根具有正实部（位于s右半平面）。")]),
    ("用routh判据判定系统稳定性",
     [("text", "den=[1,3,2,5,1];\n（构造劳斯表，检查第一列符号变化）"),
      ("img", "e3_t31_routh.png", 3.2),
      ("text", "劳斯表第一列出现2次符号变化，故有2个正实部根，系统不稳定。")]),
    ("理解给出的routh函数并用自己的方法写出routh函数",
     [("text", "重要语句说明：\nr(i,j)=(r(i-1,1)*r(i-2,j+1)-r(i-1,j+1)*r(i-2,1))/r(i-1,1); 计算劳斯表元素。\n若首列为0且该行不全为0，用eps代替；若出现全零行，则用辅助多项式求导后系数替代。"),
      ("text", "function r=routh_simple(den)\nn=length(den);\na1=den(1:2:end); a2=den(2:2:end);\nif length(a1)>length(a2), a2=[a2,0]; end\nr=[a1; a2];\nfor i=3:n\n  for j=1:length(a1)-i+2\n    r(i,j)=(r(i-1,1)*r(i-2,j+1)-r(i-1,j+1)*r(i-2,1))/r(i-1,1);\n  end\nend\nend")]),
])

para("表3-2  用MATLAB作Bode图", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令或结果"), [
    ("画 G(s)=25/(s²+4s+25) 的Bode图，并加标题",
     [("text", "sys=tf([25],[1 4 25]);bode(sys);\ntitle('Bode Diagram of G(s)=25/(s^2+4s+25)');grid on;"),
      ("img", "e3_t32_bode_a.png", 3.2)]),
    ("画 G(s)=9(s²+0.2s+1)/[s(s²+1.2s+9)] 的Bode图",
     [("text", "num=9*[1,0.2,1];\nden=conv([1,0],[1,1.2,9]);\nsys=tf(num,den);bode(sys);\ngrid on;"),
      ("img", "e3_t32_bode_b.png", 3.2)]),
])

para("表3-3  用MATLAB作Nyquist图", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令或结果"), [
    ("画 G(s)=1/(s²+0.8s+1) 的Nyquist图，并加网格和标题",
     [("text", "sys=tf([1],[1 0.8 1]);\nnyquist(sys);grid on;\ntitle('Nyquist Plot of G(s)=1/(s^2+0.8s+1)');"),
      ("img", "e3_t33_nyq_a.png", 3.0)]),
    ("画 G(s)=9(s²+0.2s+1)/[s(s²+1.2s+9)] 的Nyquist图",
     [("text", "num=9*[1,0.2,1];\nden=conv([1,0],[1,1.2,9]);\nsys=tf(num,den);nyquist(sys);grid on;"),
      ("img", "e3_t33_nyq_b.png", 3.0)]),
])

para("表3-4  用MATLAB绘制根轨迹", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
make_table(("说明", "命令或结果"), [
    ("开环传递函数 G(s)=(s+1)/(s³+4s²+2s+9)，绘制根轨迹",
     [("text", "num=[1 1];\nden=[1 4 2 9];\nrlocus(num,den);grid;\nxlabel('Real Axis');ylabel('Imaginary Axis');\ntitle('Root Locus');"),
      ("img", "e3_t34_rlocus.png", 3.0)]),
    ("绘制上行系统 K在(1,10)、步长0.5 的根轨迹图",
     [("text", "k=1:0.5:10;\nrlocus(num,den,k);grid;\ntitle('Root Locus with K from 1 to 10');"),
      ("img", "e3_t34_rlocus_k.png", 3.0)]),
])

sub("三、实验报告")
para("1．完成表3-1至3-4（见上）。")
para("2．画系统伯德图并计算相角裕量和幅值裕量。")
para("系统(a)：G(s)=500(0.0167s+1)/[s(0.05s+1)(0.0025s+1)(0.001s+1)]")
para(">>num=500*[0.0167 1];\n>>den=conv(conv([0.05 1],[0.0025 1]),conv([0.001 1],[1 0]));\n>>sys=tf(num,den);\n>>bode(sys); grid on;\n>>[Gm,Pm,Wcg,Wcp]=margin(sys)", font="Consolas")
img("e3_hw2_bode.png", width=3.8)
img("e3_hw2_margin.png", width=3.4)
para("结果：幅值裕度 Gm≈7.20（约17.1 dB），相角裕度 Pm≈45.5°，Wcg≈586.8 rad/s，Wcp≈161.7 rad/s。Pm>0 且 Gm>1，系统稳定。")
para("3．开环传递函数 G(s)=K(s²+5s+6)/(s²+8s+32)。")
para("（1）系统的根轨迹：")
para(">>num=[1 5 6]; den=[1 8 32];\n>>rlocus(num,den); grid on;\n>>xlabel('Real Axis'); ylabel('Imaginary Axis');\n>>title('Root Locus');", font="Consolas")
img("e3_hw3_rlocus.png", width=3.4)
para("开环极点 −4±4j、开环零点 −2 与 −3；根轨迹始于极点、终于零点，全部位于 s 左半平面，故任意 K>0 系统均稳定。")
para("（2）K=1 时闭环系统阶跃响应曲线：")
para(">>G=tf([1 5 6],[1 8 32]);\n>>T=feedback(1*G,1);   % K=1 单位负反馈\n>>step(T); grid on;\n>>title('K=1 闭环系统单位阶跃响应');", font="Consolas")
img("e3_hw3_step_k1.png", width=3.6)
para("K=1 时闭环传递函数 T(s)=(s²+5s+6)/(2s²+13s+38)，阶跃响应稳定收敛至稳态值（约 0.158），无明显振荡。")

doc.save(OUT)
print("saved")

